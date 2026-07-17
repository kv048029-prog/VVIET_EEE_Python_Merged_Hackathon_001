from PIL import Image

def get_position(position, img_w, img_h, obj_w, obj_h):
    margin = 20

    positions = {
        "Top Left": (margin, margin),
        "Top Right": (img_w - obj_w - margin, margin),
        "Center": ((img_w - obj_w)//2, (img_h - obj_h)//2),
        "Bottom Left": (margin, img_h - obj_h - margin),
        "Bottom Right": (img_w - obj_w - margin,
                         img_h - obj_h - margin)
    }

    return positions.get(position, positions["Bottom Right"])
  import streamlit as st
from PIL import Image, ImageDraw, ImageFont
import io

from watermark.utils import get_position


def image_page():

    st.header("🖼 Image Watermark")

    img = st.file_uploader(
        "Upload Image",
        type=["png", "jpg", "jpeg"]
    )

    if img is None:
        return

    image = Image.open(img).convert("RGBA")

    watermark = st.radio(
        "Watermark Type",
        ["Text", "Logo"]
    )

    opacity = st.slider(
        "Opacity",
        0,
        255,
        120
    )

    position = st.selectbox(
        "Position",
        [
            "Top Left",
            "Top Right",
            "Center",
            "Bottom Left",
            "Bottom Right"
        ]
    )

    overlay = Image.new("RGBA", image.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)

    if watermark == "Text":

        text = st.text_input(
            "Watermark Text",
            "AI Watermark Studio"
        )

        size = st.slider(
            "Font Size",
            15,
            80,
            35
        )

        try:
            font = ImageFont.truetype("arial.ttf", size)
        except:
            font = ImageFont.load_default()

        bbox = draw.textbbox((0, 0), text, font=font)

        w = bbox[2]-bbox[0]
        h = bbox[3]-bbox[1]

        x, y = get_position(position,
                            image.width,
                            image.height,
                            w,
                            h)

        draw.text(
            (x, y),
            text,
            font=font,
            fill=(255, 255, 255, opacity)
        )

    else:

        logo = st.file_uploader(
            "Upload Logo",
            type=["png"]
        )

        if logo is not None:

            logo = Image.open(logo).convert("RGBA")

            logo.thumbnail((150, 150))

            alpha = logo.getchannel("A").point(lambda p: opacity)
            logo.putalpha(alpha)

            x, y = get_position(position,
                                image.width,
                                image.height,
                                logo.width,
                                logo.height)

            overlay.paste(logo, (x, y), logo)

    result = Image.alpha_composite(image, overlay)

    st.image(result)

    buffer = io.BytesIO()

    result.convert("RGB").save(buffer, format="PNG")

    st.download_button(
        "Download Image",
        buffer.getvalue(),
        "watermarked.png"
    )
  import streamlit as st
import fitz


def pdf_page():

    st.header("📄 PDF Watermark")

    pdf = st.file_uploader(
        "Upload PDF",
        type=["pdf"]
    )

    if pdf is None:
        return

    text = st.text_input(
        "Watermark Text",
        "CONFIDENTIAL"
    )

    if st.button("Apply PDF Watermark"):

        doc = fitz.open(stream=pdf.read(), filetype="pdf")

        for page in doc:

            rect = page.rect

            page.insert_text(
                (rect.width/3, rect.height/2),
                text,
                fontsize=40,
                rotate=45,
                color=(0.8, 0.8, 0.8)
            )

        output = "watermarked.pdf"

        doc.save(output)

        doc.close()

        with open(output, "rb") as f:

            st.download_button(
                "Download PDF",
                f,
                output
            )
          import streamlit as st
from docx import Document
from docx.shared import Pt


def word_page():

    st.header("📝 Word Watermark")

    doc = st.file_uploader(
        "Upload DOCX",
        type=["docx"]
    )

    if doc is None:
        return

    text = st.text_input(
        "Watermark Text",
        "CONFIDENTIAL"
    )

    if st.button("Apply Word Watermark"):

        document = Document(doc)

        header = document.sections[0].header

        para = header.paragraphs[0]

        run = para.add_run(text)

        run.font.size = Pt(26)

        run.bold = True

        output = "watermarked.docx"

        document.save(output)

        with open(output, "rb") as f:

            st.download_button(
                "Download DOCX",
                f,
                output
            )
          import streamlit as st

def video_page():

    st.header("🎥 Video Watermark")

    st.info(
        "This module requires MoviePy/OpenCV.\n\n"
        "Upload a video and apply a text or logo watermark."
    )

    st.file_uploader(
        "Upload Video",
        type=["mp4", "avi", "mov"]
    )

    st.text_input(
        "Watermark Text",
        "AI Watermark Studio"
    )

    st.button("Apply Watermark")

    st.success(
        "Implement watermark processing here using MoviePy."
    )
  
